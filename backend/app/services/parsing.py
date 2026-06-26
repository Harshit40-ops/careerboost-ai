"""
parsing.py  —  LAYER 1: Clean Parsing
-------------------------------------
Turn an uploaded resume (PDF or DOCX) into clean, readable text.

Strategy:
  * PDF  -> try pdfplumber first (great at layout); if it yields too little,
            fall back to PyMuPDF (fitz).
  * DOCX -> python-docx (paragraphs + tables).

After extraction we CLEAN the text: collapse whitespace, repair broken
line-breaks, and drop junk characters. If the final text is too short we
raise ParsingError so the API can return a clear message.
"""

import io
import re

import fitz  # PyMuPDF
import pdfplumber
from docx import Document

from ..config import settings


class ParsingError(Exception):
    """Raised when we cannot extract usable text from a resume."""


# Allowed upload types -> a simple label we use internally.
SUPPORTED_TYPES = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
}


def _clean_text(text: str) -> str:
    """Normalise messy extracted text into something readable."""
    # Replace non-breaking spaces and tabs with normal spaces.
    text = text.replace("\xa0", " ").replace("\t", " ")

    # A hyphen at the end of a line usually means a word was split across
    # lines ("experi-\nence" -> "experience"). Re-join those.
    text = re.sub(r"-\n", "", text)

    # Collapse 3+ newlines into a paragraph break.
    text = re.sub(r"\n{3,}", "\n\n", text)

    # Remove control / non-printable junk but keep normal punctuation.
    text = re.sub(r"[^\x09\x0A\x0D\x20-\x7E‘-‟]", " ", text)

    # Collapse runs of spaces.
    text = re.sub(r"[ ]{2,}", " ", text)

    # Trim trailing spaces on each line.
    text = "\n".join(line.strip() for line in text.splitlines())

    return text.strip()


def _extract_pdf(data: bytes) -> str:
    """Extract text from a PDF, with a PyMuPDF fallback."""
    text = ""
    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            pages = [page.extract_text() or "" for page in pdf.pages]
            text = "\n".join(pages)
    except Exception:
        text = ""

    # If pdfplumber struggled (scanned/odd layout), try PyMuPDF.
    if len(text.strip()) < settings.MIN_RESUME_CHARS:
        try:
            with fitz.open(stream=data, filetype="pdf") as doc:
                text = "\n".join(page.get_text() for page in doc)
        except Exception as exc:  # pragma: no cover
            raise ParsingError(f"Could not read PDF: {exc}") from exc

    return text


def _extract_docx(data: bytes) -> str:
    """Extract text from a DOCX, including table cells."""
    try:
        doc = Document(io.BytesIO(data))
    except Exception as exc:
        raise ParsingError(f"Could not read DOCX: {exc}") from exc

    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def parse_resume(data: bytes, content_type: str, filename: str = "") -> str:
    """
    Main entry point.
    `data`         -> raw bytes of the uploaded file
    `content_type` -> MIME type reported by the browser
    Returns clean text, or raises ParsingError.
    """
    # Resolve the file kind from MIME type, falling back to extension.
    kind = SUPPORTED_TYPES.get(content_type)
    if kind is None:
        lower = filename.lower()
        if lower.endswith(".pdf"):
            kind = "pdf"
        elif lower.endswith(".docx"):
            kind = "docx"
        else:
            raise ParsingError("Unsupported file type. Upload a PDF or DOCX.")

    raw = _extract_pdf(data) if kind == "pdf" else _extract_docx(data)
    cleaned = _clean_text(raw)

    if len(cleaned) < settings.MIN_RESUME_CHARS:
        raise ParsingError(
            "We couldn't read enough text from this resume. "
            "If it's a scanned image, please upload a text-based PDF or DOCX."
        )
    return cleaned
